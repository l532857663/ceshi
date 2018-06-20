#!/usr/bin/env python3
#-*- coding:utf-8 -*-

import os,sys
import re
import docx
import subprocess

'''
    解析docx文件的内容函数docxAnalysis()；
.docx文件的结构分为三层:
        1、Docment对象表示整个文档；
        2、Docment包含了Paragraph对象的列表，Paragraph对象用来表示文档中的段落；
        3、一个Paragraph对象包含Run对象的列表(Word里面的文本不只是包含了字符串，还有字号、字体、颜色等等属性，都包含在style中。一个Run对象就是style相同的一段文本)
doc.paragraphs[i].text          文本段内容
doc.paragraphs[i].runs[j].text  文本段中的类型内容
'''
def docxAnalysis(fname):
    doc = docx.Document(fname)
    #print ("段落数：",len(doc.paragraphs))
    content = ""
    for i in range(len(doc.paragraphs)):
        par = doc.paragraphs[i].text
        content += par
    return (content)

        #lenRun = len(doc.paragraphs[i].runs)
        #if lenRun >= 2:
        #    print (i,"  runNum:",lenRun)
        #    for j in range(lenRun):
        #        print (doc.paragraphs[i].runs[j].text)
        #else:
        #    print (i,"  runNum:",lenRun,"   ",par)

'''
    解析doc文件的内容函数docAnalysis()；
.doc 只解析文字数据，把bytes数据转换成str数据输出
'''
def docAnalysis(fname):
    content = subprocess.check_output(["./www/data_anay/py/bin/antiword", fname])
    content = content.decode()
    return (content)
